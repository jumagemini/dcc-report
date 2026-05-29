from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from .models import DCC
import os
import base64
from django.core.files.base import ContentFile

def save_signature_from_data_url(data_url, institution):
    """Save a base64 data URL to the institution's signature field."""
    if not data_url or not data_url.startswith('data:image/png;base64,'):
        return
    # Extract the base64 part
    format, imgstr = data_url.split(';base64,')
    ext = format.split('/')[-1]
    file_name = f"sig_{institution.pk}.{ext}"
    data = ContentFile(base64.b64decode(imgstr), name=file_name)
    institution.signature.save(file_name, data, save=False)

def set_table_cant_split(table, cant_split=True):
    """Prevent a table from being broken across pages."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    for cs in tblPr.findall(qn('w:cantSplit')):
        tblPr.remove(cs)
    if cant_split:
        cant_split_elem = parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        tblPr.append(cant_split_elem)

def generate_dcc_photos_docx(dcc):
    """Generate a DOCX file with institution photos in the same layout as Excel Sheet 2."""
    doc = Document()
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    institutions = dcc.institutions.all().order_by('name')

    # Device order and column widths (in Inches) – roughly matching Excel proportions
    device_order = ['ONU', 'AP1', 'AP2', 'AP3', 'OUT']
    col_widths = [Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8)]  # uniform for simplicity

    IMG_HEIGHT = Inches(1.2)   # fixed height, width will scale proportionally

    for inst in institutions:
        # --- Institution table ---
        num_rows = 7  # name, before header, images, labels, after header, images, labels
        table = doc.add_table(rows=num_rows, cols=5)
        table.style = 'Table Grid'
        set_table_cant_split(table, True)

        # Set column widths
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

        # Row 0: Institution name (merge all columns)
        hdr_cells = table.rows[0].cells
        merged = hdr_cells[0].merge(hdr_cells[4])
        merged.text = inst.name.upper()
        for paragraph in merged.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

        # Row 1: "BEFORE INSTALLATION" header
        before_hdr = table.rows[1].cells
        merged = before_hdr[0].merge(before_hdr[4])
        merged.text = "BEFORE INSTALLATION"
        for paragraph in merged.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

        # Row 2: Before images
        before_img_row = table.rows[2]
        for col_idx, dev in enumerate(device_order):
            cell = before_img_row.cells[col_idx]
            # Remove default paragraph
            cell.paragraphs[0].clear()
            photo = inst.photos.filter(photo_type='before', device_type=dev).first()
            if photo:
                try:
                    img_path = photo.image.path
                    if os.path.exists(img_path):
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(img_path, height=IMG_HEIGHT)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        _write_na(cell)
                except Exception:
                    _write_na(cell)
            else:
                _write_na(cell)
            # Set cell vertical alignment to center
            cell.vertical_alignment = 1  # CENTER

        # Row 3: Before labels (device name + location)
        before_label_row = table.rows[3]
        for col_idx, dev in enumerate(device_order):
            cell = before_label_row.cells[col_idx]
            cell.paragraphs[0].clear()
            location = _get_location(inst, dev)
            label = f"{_device_label(dev)} {location}" if location else _device_label(dev)
            cell.paragraphs[0].text = label
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

        # Row 4: "AFTER INSTALLATION" header
        after_hdr = table.rows[4].cells
        merged = after_hdr[0].merge(after_hdr[4])
        merged.text = "AFTER INSTALLATION"
        for paragraph in merged.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

        # Row 5: After images
        after_img_row = table.rows[5]
        for col_idx, dev in enumerate(device_order):
            cell = after_img_row.cells[col_idx]
            cell.paragraphs[0].clear()
            photo = inst.photos.filter(photo_type='after', device_type=dev).first()
            if photo:
                try:
                    img_path = photo.image.path
                    if os.path.exists(img_path):
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(img_path, height=IMG_HEIGHT)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        _write_na(cell)
                except Exception:
                    _write_na(cell)
            else:
                _write_na(cell)
            cell.vertical_alignment = 1

        # Row 6: After labels
        after_label_row = table.rows[6]
        for col_idx, dev in enumerate(device_order):
            cell = after_label_row.cells[col_idx]
            cell.paragraphs[0].clear()
            location = _get_location(inst, dev)
            label = f"{_device_label(dev)} {location}" if location else _device_label(dev)
            cell.paragraphs[0].text = label
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

        # Add a blank paragraph after each institution table
        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _write_na(cell):
    """Write 'N/A' centered in a cell."""
    cell.paragraphs[0].text = 'N/A'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.font.size = Pt(11)
        run.bold = True

def _device_label(dev):
    """Return friendly device name."""
    return {
        'ONU': 'ONU',
        'AP1': 'INDOOR AP1',
        'AP2': 'INDOOR AP2',
        'AP3': 'INDOOR AP3',
        'OUT': 'OUTDOOR AP1'
    }.get(dev, dev)

def _get_location(inst, dev):
    """Return installation location for a device."""
    return {
        'ONU': inst.onu_location,
        'AP1': inst.indoor_ap1_location,
        'AP2': inst.indoor_ap2_location,
        'AP3': inst.indoor_ap3_location,
        'OUT': inst.outdoor_ap_location,
    }.get(dev, '')        

def add_institution_block(cell, inst):
    # Clear default paragraph
    cell.paragraphs[0].clear()

    devices = []
    if inst.indoor_ap1_serial:
        devices.append(('INDOOR AP1', inst.indoor_ap1_location))
    if inst.indoor_ap2_serial:
        devices.append(('INDOOR AP2', inst.indoor_ap2_location))
    if inst.indoor_ap3_serial:
        devices.append(('INDOOR AP3', inst.indoor_ap3_location))
    if inst.outdoor_ap_serial:
        devices.append(('OUTDOOR AP1', inst.outdoor_ap_location))

    if not devices:
        p = cell.add_paragraph()
        run = p.add_run(inst.name.upper())
        run.bold = True
        run.font.size = Pt(10)
        return

    inner_table = cell.add_table(rows=1+len(devices), cols=2)
    set_table_cant_split(inner_table, True)
    inner_table.style = 'Table Grid'

    for row in inner_table.rows:
        row.height = Inches(0.35)

    # Institution name row (merged)
    hdr_cells = inner_table.rows[0].cells
    merged = hdr_cells[0].merge(hdr_cells[1])
    merged.text = inst.name.upper()
    for paragraph in merged.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(11)

    for idx, (dev_name, location) in enumerate(devices, start=1):
        row_cells = inner_table.rows[idx].cells
        row_cells[0].text = dev_name
        row_cells[1].text = location
        for paragraph in row_cells[0].paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        for paragraph in row_cells[1].paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # No spacer paragraph added

def generate_dcc_stickers_docx(dcc):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    institutions = list(dcc.institutions.all().order_by('name'))
    chunk_size = 8
    for i in range(0, len(institutions), chunk_size):
        chunk = institutions[i:i+chunk_size]
        if i > 0:
            doc.add_page_break()

        master = doc.add_table(rows=4, cols=2)
        master.autofit = True
        set_table_cant_split(master, True)

        # Shrink master table spacing
        for row in master.rows:
            row.height = Inches(0.01)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                # Remove cell margins
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for el in tcPr.findall(qn('w:tcMar')):
                    tcPr.remove(el)

        for row_idx in range(4):
            for col_idx in range(2):
                inst_idx = row_idx * 2 + col_idx
                if inst_idx < len(chunk):
                    cell = master.cell(row_idx, col_idx)
                    add_institution_block(cell, chunk[inst_idx])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer